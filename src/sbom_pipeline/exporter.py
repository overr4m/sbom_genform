"""
Экспорт SBOM-данных в отчёты:
  • Excel (.xlsx) — Лист 1: Компоненты, Лист 2: Уязвимости
  • Word  (.docx) — python-docx
  • ODT   (.odt)  — odfpy
"""

from __future__ import annotations

import hashlib
import logging
import os
from pathlib import Path
from typing import Any, Dict, List, Optional

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from odf.opendocument import OpenDocumentText
from odf.style import Style, TableCellProperties, TextProperties, ParagraphProperties
from odf.table import Table, TableRow, TableCell
from odf.text import P

from .vuln_merger import VulnFinding

# Колонки листа "Компоненты"
_COMP_COLUMNS = [
    "№ п/п",
    "Наименование компонента",
    "Версия компонента",
    "Тип пакета / тип компонента",
    "PURL / технический идентификатор компонента",
    "Язык (языки)",
    "Признак принадлежности к поверхности атаки",
    "Признак выполнения функций безопасности",
    "Принадлежность к контейнерному образу",
    "Роль компонента в составе контейнерного образа",
    "Адрес веб-ресурса",
]

# Колонки листа "Уязвимости"
_VULN_COLUMNS = [
    "Компонент",
    "Версия",
    "CVE / ID",
    "CVSS",
    "Критичность",
    "Описание",
    "Сканер",
    "Исправлено в версии",
    "Рекомендация / компенсирующая мера",
    "Статус допустимости в рассматриваемой конфигурации",
]
_BDU_VULN_COLUMN = "BDU / ID"
_SEVERITY_COLUMN = "Критичность"

# Цвета критичности для Word
_SEVERITY_COLORS: Dict[str, RGBColor] = {
    "CRITICAL": RGBColor(0xC0, 0x00, 0x00),
    "HIGH":     RGBColor(0xFF, 0x00, 0x00),
    "MEDIUM":   RGBColor(0xFF, 0x82, 0x00),
    "LOW":      RGBColor(0xFF, 0xC0, 0x00),
    "UNKNOWN":  RGBColor(0x80, 0x80, 0x80),
}


class Exporter:
    def __init__(
        self,
        dependencies: list,
        vulns: Optional[List[VulnFinding]] = None,
        sbom_path: Optional[str] = None,
        include_bdu: bool = False,
    ) -> None:
        self.deps = dependencies
        self.vulns: List[VulnFinding] = vulns or []
        self.sbom_path = sbom_path
        self.include_bdu = include_bdu
        logging.info(
            f"Exporter: {len(self.deps)} зависимостей, {len(self.vulns)} уязвимостей"
        )

    # ------------------------------------------------------------------
    # Excel
    # ------------------------------------------------------------------

    def exportToExcel(self, report: str) -> None:
        path = Path(report)
        path.parent.mkdir(parents=True, exist_ok=True)
        try:
            vuln_columns = self._vuln_columns()
            with pd.ExcelWriter(str(path), engine="openpyxl") as writer:
                # Лист 1: Компоненты
                comp_df = pd.DataFrame(self._comp_rows(), columns=_COMP_COLUMNS)
                comp_df.to_excel(writer, index=False, sheet_name="Компоненты")

                # Лист 2: Уязвимости
                vuln_df = pd.DataFrame(self._vuln_rows(), columns=vuln_columns)
                vuln_df.to_excel(writer, index=False, sheet_name="Уязвимости")

            self._write_sig(path)
            logging.info(f"[exporter] Excel → {path}")
        except Exception as e:
            logging.exception(f"[exporter] Excel ошибка: {e}")

    # ------------------------------------------------------------------
    # Word (.docx)
    # ------------------------------------------------------------------

    def exportToDocx(self, report: str) -> None:
        path = Path(report)
        path.parent.mkdir(parents=True, exist_ok=True)
        try:
            doc = Document()
            doc.add_heading("Отчёт по компонентам SBOM", level=1)

            # --- Таблица компонентов ---
            doc.add_heading("Компоненты", level=2)
            comp_rows = self._comp_rows()
            t = doc.add_table(rows=1 + len(comp_rows), cols=len(_COMP_COLUMNS))
            t.style = "Table Grid"
            # Заголовок
            for i, col in enumerate(_COMP_COLUMNS):
                cell = t.rows[0].cells[i]
                cell.text = col
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.size = Pt(9)
            # Данные
            for r_idx, row in enumerate(comp_rows, start=1):
                for c_idx, val in enumerate(row.values()):
                    t.rows[r_idx].cells[c_idx].text = str(val)

            # --- Таблица уязвимостей ---
            if self.vulns:
                doc.add_heading("Уязвимости", level=2)
                vuln_columns = self._vuln_columns()
                vuln_rows = self._vuln_rows()
                vt = doc.add_table(rows=1 + len(vuln_rows), cols=len(vuln_columns))
                vt.style = "Table Grid"
                for i, col in enumerate(vuln_columns):
                    cell = vt.rows[0].cells[i]
                    cell.text = col
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                        run.font.size = Pt(9)
                for r_idx, row in enumerate(vuln_rows, start=1):
                    for c_idx, val in enumerate(row.values()):
                        cell = vt.rows[r_idx].cells[c_idx]
                        cell.text = str(val)
                        # Подсветка критичности
                        if vuln_columns[c_idx] == _SEVERITY_COLUMN:
                            color = _SEVERITY_COLORS.get(str(val).upper(), _SEVERITY_COLORS["UNKNOWN"])
                            for run in cell.paragraphs[0].runs:
                                run.font.color.rgb = color
                                run.bold = True

            doc.save(str(path))
            self._write_sig(path)
            logging.info(f"[exporter] Word → {path}")
        except Exception as e:
            logging.exception(f"[exporter] Word ошибка: {e}")

    # ------------------------------------------------------------------
    # ODT
    # ------------------------------------------------------------------

    def exportToOdt(self, report: str) -> None:
        path = Path(report)
        path.parent.mkdir(parents=True, exist_ok=True)
        try:
            doc = OpenDocumentText()

            title_style = Style(name="Title", family="paragraph")
            title_style.addElement(ParagraphProperties())
            title_style.addElement(TextProperties(fontsize="16pt", fontweight="bold"))
            doc.styles.addElement(title_style)

            cell_style = Style(name="CellBorders", family="table-cell")
            cell_style.addElement(TableCellProperties(border="0.05pt solid #808080"))
            doc.styles.addElement(cell_style)

            # --- Таблица компонентов ---
            doc.text.addElement(P(text="Компоненты", stylename=title_style))
            doc.text.addElement(self._make_odt_table("SBOM_Components", _COMP_COLUMNS, self._comp_rows(), cell_style))

            # --- Таблица уязвимостей ---
            if self.vulns:
                doc.text.addElement(P(text="Уязвимости", stylename=title_style))
                doc.text.addElement(
                    self._make_odt_table(
                        "SBOM_Vulns",
                        self._vuln_columns(),
                        self._vuln_rows(),
                        cell_style,
                    )
                )

            doc.save(str(path))
            self._write_sig(path)
            logging.info(f"[exporter] ODT → {path}")
        except Exception as e:
            logging.exception(f"[exporter] ODT ошибка: {e}")

    # ------------------------------------------------------------------
    # Вспомогательные методы
    # ------------------------------------------------------------------

    def _comp_rows(self) -> List[Dict[str, Any]]:
        rows = []
        for i, dep in enumerate(self.deps, start=1):
            dep_types: list = [
                t for t in (getattr(dep, "dep_type", None) or getattr(dep, "depType", []) or [])
                if isinstance(t, str)
            ]
            attack_surface = ", ".join(
                t for t in dep_types if "attack" in t.lower() or "поверхность" in t.lower()
            ) or getattr(dep, "attack_surface", "")
            security_func = ", ".join(
                t for t in dep_types if "security" in t.lower() or "безопасност" in t.lower()
            ) or getattr(dep, "security_function", "")
            rows.append({
                _COMP_COLUMNS[0]: i,
                _COMP_COLUMNS[1]: getattr(dep, "name", ""),
                _COMP_COLUMNS[2]: getattr(dep, "version", ""),
                _COMP_COLUMNS[3]: getattr(dep, "package_type", "") or getattr(dep, "component_type", ""),
                _COMP_COLUMNS[4]: getattr(dep, "purl", ""),
                _COMP_COLUMNS[5]: ", ".join(
                    getattr(dep, "src_langs", None) or getattr(dep, "srcLangs", []) or []
                ),
                _COMP_COLUMNS[6]: attack_surface,
                _COMP_COLUMNS[7]: security_func,
                _COMP_COLUMNS[8]: getattr(dep, "container_image", ""),
                _COMP_COLUMNS[9]: getattr(dep, "container_role", ""),
                _COMP_COLUMNS[10]: getattr(dep, "source", None) or "",
            })
        return rows

    def _vuln_columns(self) -> List[str]:
        columns = list(_VULN_COLUMNS)
        if self.include_bdu:
            columns.insert(3, _BDU_VULN_COLUMN)
        return columns

    def _vuln_rows(self) -> List[Dict[str, Any]]:
        rows: List[Dict[str, Any]] = []
        for v in self.vulns:
            row: Dict[str, Any] = {
                _VULN_COLUMNS[0]: v.component_name,
                _VULN_COLUMNS[1]: v.component_version,
                _VULN_COLUMNS[2]: v.cve_id,
                _VULN_COLUMNS[3]: v.cvss_score,
                _VULN_COLUMNS[4]: v.severity_upper,
                _VULN_COLUMNS[5]: v.description[:200] if v.description else "",
                _VULN_COLUMNS[6]: v.scanner,
                _VULN_COLUMNS[7]: v.fixed_version,
                _VULN_COLUMNS[8]: getattr(v, "recommendation", ""),
                _VULN_COLUMNS[9]: getattr(v, "acceptability_status", ""),
            }
            if self.include_bdu:
                items = list(row.items())
                items.insert(3, (_BDU_VULN_COLUMN, v.bdu_id or ""))
                row = dict(items)
            rows.append(row)
        return rows

    @staticmethod
    def _make_odt_table(
        name: str,
        columns: List[str],
        rows: List[Dict[str, Any]],
        cell_style: Any,
    ) -> Table:
        table = Table(name=name)
        # Заголовок
        hrow = TableRow()
        for col in columns:
            cell = TableCell(stylename=cell_style)
            cell.addElement(P(text=col))
            hrow.addElement(cell)
        table.addElement(hrow)
        # Данные
        for row in rows:
            trow = TableRow()
            for val in row.values():
                cell = TableCell(stylename=cell_style)
                cell.addElement(P(text=str(val)))
                trow.addElement(cell)
            table.addElement(trow)
        return table

    @staticmethod
    def _write_sig(file_path: Path) -> None:
        """Записать SHA-256 контрольную сумму рядом с файлом."""
        try:
            with open(file_path, "rb") as f:
                digest = hashlib.sha256(f.read()).hexdigest()
            sig_path = file_path.with_suffix(file_path.suffix + ".sig")
            sig_path.write_text(f"SHA256={digest}\n", encoding="utf-8")
            logging.info(f"[exporter] Подпись → {sig_path}")
        except Exception as e:
            logging.warning(f"[exporter] Не удалось создать .sig: {e}")
